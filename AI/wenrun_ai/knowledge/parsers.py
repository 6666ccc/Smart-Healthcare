from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from .types import KnowledgeDocumentError


@dataclass(frozen=True)
class ParsedDocument:
    text: str


def parse_document(file_bytes: bytes, file_name: str) -> ParsedDocument:
    if not file_bytes:
        raise KnowledgeDocumentError("文档内容为空")

    extension = Path(file_name).suffix.lower()
    if extension in {".txt", ".md"}:
        text = _parse_text(file_bytes)
    elif extension == ".docx":
        text = _parse_docx(file_bytes)
    elif extension == ".pdf":
        text = _parse_pdf(file_bytes)
    else:
        raise KnowledgeDocumentError("仅支持 PDF、DOCX、TXT 和 Markdown 文件")

    normalized = _normalize(text)
    if not normalized:
        raise KnowledgeDocumentError("文档没有可用文本；扫描 PDF 暂不支持")
    return ParsedDocument(text=normalized)


def _parse_text(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise KnowledgeDocumentError("文本文件必须使用 UTF-8 编码") from exc


def _parse_docx(file_bytes: bytes) -> str:
    try:
        document = Document(BytesIO(file_bytes))
        return "\n\n".join(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
    except Exception as exc:
        raise KnowledgeDocumentError("DOCX 文档损坏或无法解析") from exc


def _parse_pdf(file_bytes: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(file_bytes))
        if reader.is_encrypted:
            raise KnowledgeDocumentError("不支持加密 PDF")
        return "\n\n".join((page.extract_text() or "").strip() for page in reader.pages)
    except KnowledgeDocumentError:
        raise
    except Exception as exc:
        raise KnowledgeDocumentError("PDF 文档损坏或无法解析") from exc


def _normalize(text: str) -> str:
    lines = [line.rstrip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    normalized: list[str] = []
    previous_blank = False
    for line in lines:
        blank = not line.strip()
        if blank and previous_blank:
            continue
        normalized.append("" if blank else line.strip())
        previous_blank = blank
    return "\n".join(normalized).strip()
